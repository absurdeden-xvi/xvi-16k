from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("/Users/mymac/Documents/Codex/2026-06-16/project/outputs/XVI-十六开-产品与视觉设计问卷.docx")
FONT = "SimSun"
ACCENT = "FF3B18"
INK = "171717"
MUTED = "666666"
LIGHT = "F3F3F0"


SECTIONS = [
    ("一、产品本质", [
        ("用一句话描述，你希望 XVI 帮用户完成什么？", [], 3, True),
        ("它更接近哪种产品？", ["专业排版工具", "面向普通人的轻量创作工具", "社交媒体长图生成器", "写作者专属发布工具", "视觉作品生成器"], 1, True),
        ("用户使用 XVI 时，最重要的感受是什么？", ["快速", "自由", "专业", "有灵感", "安静专注", "结果惊艳"], 1, False),
        ("它应该更像“工具”，还是更像“创作空间”？为什么？", [], 2, False),
        ("用户为什么不会直接使用 Canva、Figma、截图或其他长图工具？", [], 3, False),
        ("最终成品更应该突出什么？", ["文字内容", "排版品味", "作者身份", "XVI 的品牌辨识度"], 1, False),
        ("你希望用户如何向别人介绍它？", [], 2, False),
    ]),
    ("二、目标用户", [
        ("第一批最核心的用户是谁？", ["小说作者", "同人作者", "随笔作者", "诗歌作者", "社交媒体创作者", "编辑与出版从业者", "普通记录者"], 1, True),
        ("他们通常在哪里发布长图？", [], 2, False),
        ("他们一次处理的文本通常有多少字？", [], 1, False),
        ("他们是否了解字号、行距、字距、网格等排版概念？", [], 2, False),
        ("他们更需要自动完成，还是精细控制？", [], 2, False),
        ("手机端和电脑端的重要程度分别是多少？请按 1-5 评分。", [], 1, False),
        ("用户愿意花多久完成一次排版？", [], 1, False),
        ("用户最容易在哪一步失去耐心？", [], 2, False),
    ]),
    ("三、使用流程", [
        ("你理想中的主要流程是哪一种？", ["输入文本 → 自动生成 → 微调 → 导出", "输入文本 → 选择模板 → 自动生成 → 导出", "输入文本 → AI 分析内容 → 推荐版式 → 导出", "像专业编辑器一样边写边排"], 1, True),
        ("标题、署名、正文是否应该分开输入？", [], 1, False),
        ("是否需要副标题、日期、章节名、题记？", [], 2, False),
        ("智能分段应该如何触发？", ["默认自动执行", "生成前展示建议", "只在用户主动选择时执行"], 1, False),
        ("自动排版后，用户修改正文时应该怎样处理？", ["自动重新排版", "提示重新生成", "保留当前版式，仅替换文字"], 1, False),
        ("是否需要撤销与重做？", [], 1, False),
        ("是否需要比较不同排版方案？", [], 1, False),
        ("是否需要保存草稿和历史版本？", [], 1, False),
        ("最理想的一次操作需要多少次点击？", [], 1, False),
    ]),
    ("四、自动排版", [
        ("你认为“自动排版”至少应该自动决定哪些内容？", ["字号", "行距", "字距", "段距", "页边距", "标题比例", "分段", "对齐方式", "画布宽度", "配色"], 1, True),
        ("自动排版应该主要依据什么？", ["总字数", "文体", "情绪", "发布平台", "用户选择的模板"], 1, False),
        ("是否应该识别小说、诗歌、对白、书信、随笔等文本结构？", [], 2, False),
        ("诗歌是否必须保留原始换行？", [], 1, False),
        ("小说对白是否需要特殊缩进或段距？", [], 1, False),
        ("用户是否应该看到系统做了哪些调整？", [], 1, False),
        ("你更偏好一个最佳结果，还是一次生成 3 个候选方案？", [], 2, False),
        ("自动排版失败时，怎样才算“失败”？", [], 3, False),
    ]),
    ("五、视觉气质", [
        ("用 3-5 个词描述理想的 XVI。", [], 2, True),
        ("用 3-5 个词描述它绝对不能呈现的感觉。", [], 2, True),
        ("它应该偏向哪种视觉方向？", ["瑞士国际主义", "现代编辑出版", "时装杂志", "实验性平面设计", "极简软件工具", "数字原生界面", "当代中文排版"], 1, True),
        ("“高级感”对你具体意味着什么？", ["留白", "字体品质", "精确网格", "微妙动效", "材质感", "克制配色", "不对称构图"], 2, True),
        ("UI 应该更冷静理性，还是更有情绪和表达？", [], 2, False),
        ("是否接受大胆颜色、大字号和非对称布局？", [], 1, False),
        ("是否接受黑色或深色主界面？", [], 1, False),
        ("你喜欢哪些网站、杂志、软件或品牌的设计？", [], 3, False),
        ("有哪些设计是你明确讨厌的？", [], 3, False),
        ("你希望 UI 本身抢眼，还是退后服务于作品？", [], 2, False),
    ]),
    ("六、品牌与“16”", [
        ("产品主名称应该如何显示？", ["XVI", "十六开", "XVI / 十六开", "十六开 XVI"], 1, True),
        ("“16”应该明显存在，还是成为隐藏设计规则？", [], 2, False),
        ("哪些地方适合出现 16？", ["Logo", "网格", "间距", "模板编号", "导出标记", "动效"], 1, False),
        ("导出的图片是否应该默认带 XVI 标记？", [], 1, False),
        ("“十六开”的纸张与出版含义是否需要被用户理解？", [], 2, False),
        ("品牌更应该面向中文用户，还是天然国际化？", [], 2, False),
        ("xvi.page、xvi.ink、xvi.studio 中你偏好哪种气质？", [], 2, False),
    ]),
    ("七、编辑器界面", [
        ("你更喜欢哪种界面结构？", ["左侧设置、右侧预览", "先输入，再进入独立预览页面", "顶部工具栏、中央画布", "全屏写作，生成后切换到排版工作台"], 1, True),
        ("设置项应该全部展开，还是按“基础 / 高级”分层？", [], 2, False),
        ("滑杆和数字输入哪一种更常用？", [], 1, False),
        ("模板应该如何显示？", ["色块", "小型成品缩略图", "大幅作品预览"], 1, False),
        ("是否需要隐藏所有控制，只查看最终作品？", [], 1, False),
        ("预览区应该展示完整长图，还是分页 / 分段浏览？", [], 2, False),
        ("是否需要画布导航缩略图？", [], 1, False),
        ("是否需要键盘快捷键？", [], 1, False),
        ("空白状态应该安静，还是主动提供示例和灵感？", [], 2, False),
    ]),
    ("八、模板与风格", [
        ("模板应该改变什么？", ["仅颜色", "字体和颜色", "完整版式结构", "标题、正文、署名的全部关系"], 1, True),
        ("第一版最需要哪些模板类型？", ["小说", "随笔", "诗歌", "书信", "摘录", "宣言", "极简", "实验排版"], 1, False),
        ("模板应该有情绪名称，还是编号和专业名称？", [], 2, False),
        ("用户能否保存自己的模板？", [], 1, False),
        ("是否允许模板使用装饰图形、边框、图片和纹理？", [], 2, False),
        ("配色应该使用固定方案，还是开放完整调色系统？", [], 2, False),
        ("是否需要随机探索或“换一个方案”？", [], 1, False),
        ("风格选择应该发生在输入前还是生成后？", [], 2, False),
    ]),
    ("九、图片与导出", [
        ("主要发布平台有哪些？", [], 2, True),
        ("应该导出单张长图，还是自动拆分多张？", [], 2, False),
        ("是否需要平台尺寸预设？", [], 1, False),
        ("最大文本长度大约是多少？", [], 1, False),
        ("是否需要导出高清倍图？", [], 1, False),
        ("是否需要 JPG、PNG、PDF？", [], 1, False),
        ("是否需要透明背景？", [], 1, False),
        ("是否需要封面和正文分别导出？", [], 1, False),
        ("用户是否应该控制文件大小和压缩质量？", [], 1, False),
        ("是否需要添加作者水印、账号或二维码？", [], 2, False),
    ]),
    ("十、优先级与边界", [
        ("V1 最不可妥协的三个能力是什么？", [], 3, True),
        ("当前版本最需要删除或弱化的三个部分是什么？", [], 3, True),
        ("当前版本哪些地方已经可以保留？", [], 3, False),
        ("你更愿意先做到哪一种？", ["功能完整但设计普通", "功能克制但视觉和体验非常成熟"], 1, False),
        ("哪些功能看似有用，但会让产品变得臃肿？", [], 3, False),
        ("是否考虑未来加入 AI 文本理解或风格推荐？", [], 2, False),
        ("是否考虑登录、云端草稿和跨设备同步？", [], 2, False),
        ("是否计划开源？", [], 1, False),
        ("最终是否考虑收费？哪些能力可以收费？", [], 3, False),
        ("三个月后，什么结果能让你认为这个项目成功？", [], 3, False),
    ]),
]


def set_run_font(run, size=10.5, bold=False, color=INK):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, 9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    tail = paragraph.add_run(" 页")
    set_run_font(tail, 9, color=MUTED)


def add_response(doc, lines):
    p = doc.add_paragraph()
    p.style = "Response"
    run = p.add_run("回答：")
    set_run_font(run, 10.5, color=MUTED)
    for _ in range(max(1, lines)):
        run.add_break()
    shade_paragraph(p, LIGHT)
    return p


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, before, after in [
        ("Heading 1", 16, 18, 10),
        ("Heading 2", 13, 14, 7),
        ("Heading 3", 11.5, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(INK)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    question = styles.add_style("Question", 1)
    question.font.name = FONT
    question._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    question.font.size = Pt(10.5)
    question.font.bold = True
    question.paragraph_format.space_before = Pt(7)
    question.paragraph_format.space_after = Pt(4)
    question.paragraph_format.keep_with_next = True

    option = styles.add_style("Option", 1)
    option.font.name = FONT
    option._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    option.font.size = Pt(10.5)
    option.paragraph_format.left_indent = Inches(0.2)
    option.paragraph_format.space_after = Pt(2)
    option.paragraph_format.keep_with_next = True

    response = styles.add_style("Response", 1)
    response.font.name = FONT
    response._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    response.font.size = Pt(10.5)
    response.paragraph_format.left_indent = Inches(0.08)
    response.paragraph_format.right_indent = Inches(0.08)
    response.paragraph_format.space_before = Pt(3)
    response.paragraph_format.space_after = Pt(8)
    response.paragraph_format.line_spacing = 1.25

    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = header_p.add_run("XVI / 十六开  ·  产品与视觉设计问卷")
    set_run_font(hr, 9, bold=True, color=MUTED)
    add_page_number(section.footer.paragraphs[0])

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(8)
    kr = kicker.add_run("XVI / PRODUCT DISCOVERY / 016")
    set_run_font(kr, 9, bold=True, color=ACCENT)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(6)
    tr = title.add_run("XVI / 十六开\n产品与视觉设计问卷")
    set_run_font(tr, 24, bold=True, color=INK)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    sr = subtitle.add_run("用于明确产品定位、核心流程、自动排版逻辑与下一轮 UI 重构方向")
    set_run_font(sr, 10.5, color=MUTED)

    instruction = doc.add_paragraph()
    instruction.paragraph_format.space_after = Pt(12)
    shade_paragraph(instruction, LIGHT)
    ir = instruction.add_run(
        "填写说明：可直接在浅灰色“回答”区域输入。选择题可将“□”改为“■”或在选项后补充说明。"
        "不确定的问题可以填写“暂不确定”。标有“★”的问题是第一轮 UI 重构的核心依据。"
    )
    set_run_font(ir, 10.5, color=INK)

    metadata = doc.add_paragraph()
    metadata.paragraph_format.space_after = Pt(10)
    mr = metadata.add_run("填写人：____________________    日期：____________________")
    set_run_font(mr, 10.5, color=MUTED)

    question_number = 1
    for section_title, questions in SECTIONS:
        h = doc.add_paragraph(section_title, style="Heading 1")
        shade_paragraph(h, "E8E8E4")
        for text, choices, response_lines, core in questions:
            p = doc.add_paragraph(style="Question")
            marker = "★ " if core else ""
            r = p.add_run(f"{question_number}. {marker}{text}")
            set_run_font(r, 10.5, bold=True, color=INK)
            for choice in choices:
                op = doc.add_paragraph(style="Option")
                rr = op.add_run(f"□ {choice}")
                set_run_font(rr, 10.5, color=INK)
            add_response(doc, response_lines)
            question_number += 1

    doc.core_properties.title = "XVI / 十六开 - 产品与视觉设计问卷"
    doc.core_properties.subject = "产品定位与 UI 重构调研"
    doc.core_properties.author = "XVI / 十六开"
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_document()
